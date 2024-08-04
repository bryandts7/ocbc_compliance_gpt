import fitz
from docx import Document as DocxDocument
from openpyxl import load_workbook
from langchain_core.documents import Document
import pandas as pd
from paddleocr import PaddleOCR
import os
import xlrd
import datetime

# ==================== EXTRACT TEXT FROM PDF ====================
def extract_text_and_images_from_page(doc, page, ocr, treshold):
    text = page.get_text()
    image_text = ""
    image_list = page.get_images(full=True)
    # Iterate through all images found on the page
    for image_info in image_list:
        xref = image_info[0]
        image_dict = doc.extract_image(xref)
        image_bytes = image_dict['image']

        if image_bytes is not None:
            # Use PaddleOCR to extract text from the image
            ocr_result = ocr.ocr(image_bytes)
            # Check if OCR result is valid before processing
            if ocr_result and ocr_result != [None]:
                for result in ocr_result:
                    for res in result:
                        text_tuple = res[1]
                        text_string = text_tuple[0]
                        # For confidence threshold
                        text_confidence = text_tuple[1]
                        if text_confidence > treshold:
                            image_text += text_string + '\n'
    # Combine page text and image text
    return text + "\n" + image_text


def extract_text_from_pdf(file_path, ocr, treshold):
    # Load the PDF file
    doc = fitz.open(file_path)
    # text = ""
    text_with_page = []
    # Iterate through all pages in the PDF
    for page_num in range(len(doc)):
        page = doc[page_num]
        # Extract text and images from the page
        page_text = extract_text_and_images_from_page(doc, page, ocr, treshold)
        # text += page_text + "\n"
        # text += f"page={page_num + 1}\n====================\n{page_text}\n====================\n"
        text_with_page.append({
            "page_number": page_num + 1,
            "text": page_text
        })
    return text_with_page


# ==================== EXTRACT TEXT FROM DOCX ====================

def extract_text_from_docx(file_path):
    doc = DocxDocument(file_path)
    text = ''
    for para in doc.paragraphs:
        text += para.text + '\n'
    return text


# ==================== EXTRACT TEXT FROM EXCEL ====================

def extract_text_from_excel(excel_path):
    if excel_path.endswith('.xls'):
        # Use xlrd for .xls files
        wb = xlrd.open_workbook(excel_path)
        text_data = []
        for sheet in wb.sheets():
            for row in range(sheet.nrows):
                for col in range(sheet.ncols):
                    cell_value = sheet.cell_value(row, col)
                    if cell_value:
                        text_data.append(str(cell_value))
    else:
        # Use openpyxl for .xlsx and .xlsm files
        wb = load_workbook(excel_path)
        text_data = []
        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                for cell in row:
                    if cell is not None:
                        text_data.append(str(cell))

    return ' '.join(text_data)


# # ==================== MAIN ====================

def convert_text_to_document(metadata, text, page_number=None):
    metadata = extract_metadata_from_dataframe(metadata)
    metadata.pop('effective_date_id_str')
    if page_number is None:
        metadata['page_number'] = None
        return Document(
            page_content=text,
            metadata=metadata
        )
    else:
        page_content = text
        metadata['page_number'] = page_number
        return Document(
            page_content=page_content,
            metadata=metadata
        )


def extract_metadata_from_dataframe(metadata):
    # Define the mapping for Indonesian month names to their numerical values
    month_mapping = {
        'Januari': 1, 'Februari': 2, 'Maret': 3, 'April': 4, 'Mei': 5,
        'Juni': 6, 'Juli': 7, 'Agustus': 8, 'September': 9, 'Oktober': 10,
        'November': 11, 'Desember': 12
    }
    
    # Convert 'NO DATA' to None
    def convert_no_data(value):
        return None if value == 'NO DATA' else value
    
    # Convert metadata values
    metadata = {
        "doc_id": convert_no_data(metadata['doc_id']),
        "title": convert_no_data(metadata['title']),
        "sector": convert_no_data(metadata['sektor']),
        "subsector": convert_no_data(metadata['subsektor']),
        "regulation_type": convert_no_data(metadata['jenis_regulasi']),
        "regulation_number": convert_no_data(metadata['nomor_regulasi']),
        "effective_date": convert_no_data(metadata['tanggal_berlaku']),
        "effective_date_id_str": convert_no_data(metadata['tanggal_berlaku']),
        "file_url": convert_no_data(metadata['file_url']),
    }
    
    # Extract and convert the effective_date to datetime
    if metadata['effective_date'] is not None:
        date_str = metadata['effective_date']
        day, month_name, year = date_str.split()
        month = month_mapping[month_name]
        effective_date = datetime.datetime(int(year), month, int(day)).strftime('%Y/%m/%d')
        metadata['effective_date'] = effective_date
    else:
        metadata['effective_date'] = None

    return metadata

# Format metadata to string
def format_metadata(metadata):
    metadata.pop('doc_id')
    metadata.pop('file_url')
    metadata.pop('effective_date')
    metadata['effective_date'] = metadata.pop('effective_date_id_str')
    if 'page_number' in metadata:
        metadata.pop('page_number')

    str_metadata = f"METADATA:\n"
    for key, value in metadata.items():
        str_metadata += f"{key}: {value}\n"
    str_metadata += '-'*10 + '\n\n'
    return str_metadata


def extract_all_documents_in_directory(documents_dir, metadata_path, treshold=0.98):
    ocr = PaddleOCR(use_angle_cls=True, lang='id', show_log=False, use_gpu=True)

    docs = []

    df_metadata = pd.read_csv(metadata_path)
    i = 0
    for file in os.listdir(documents_dir):
        i += 1
        file_path = os.path.join(documents_dir, file)
        file_metadata = df_metadata[df_metadata['new_filename']
                                    == file].iloc[0]
        if file.endswith('.pdf'):
            metadata = extract_metadata_from_dataframe(file_metadata)
            text_with_page = extract_text_from_pdf(
                file_path, ocr, treshold)
            for page in text_with_page:
                metadata_with_page = metadata.copy()
                metadata_with_page['page_number'] = page['page_number']
                metadata_str = format_metadata(metadata_with_page)
                document = convert_text_to_document(
                    text=metadata_str + page['text'], metadata=file_metadata, page_number=page['page_number'])
                docs.append(document)
        elif file.endswith('.xlsm') or file.endswith('.xlsx') or file.endswith('.xls'):
            metadata = extract_metadata_from_dataframe(file_metadata)
            metadata_str = format_metadata(metadata)
            text = metadata_str + extract_text_from_excel(file_path) + '\n'
            document = convert_text_to_document(
                text=text, metadata=file_metadata)
            docs.append(document)
        elif file.endswith('.docx'):
            metadata = extract_metadata_from_dataframe(file_metadata)
            metadata_str = format_metadata(metadata)
            text = metadata_str + extract_text_from_docx(file_path) + '\n'
            document = convert_text_to_document(
                text=text, metadata=file_metadata)
            docs.append(document)
        print(f"{i}: {file}") 
    print(f"Read {len(docs)} documents")
    return docs
